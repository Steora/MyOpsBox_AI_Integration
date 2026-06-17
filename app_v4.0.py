# --- UPDATE THE DOCX IMPORTS AT THE TOP OF YOUR APP.PY ---
import streamlit as st
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor  # <-- Added Inches here!
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from datetime import datetime, timedelta
from openai import OpenAI
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import io

# Initialize the modern client object safely using your Streamlit secrets ecosystem
ai_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
fathom_key = st.secrets["FATHOM_API_KEY"]

def clean_ai_markdown(text):
    """Removes raw markdown code blocks from the AI string."""
    if text.startswith("```markdown"):
        text = text[11:]
    elif text.startswith("```"):
        text = text[3:]
        
    if text.endswith("```"):
        text = text[:-3]
        
    return text.strip()

# --- CONFIGURATION & CONSTANTS ---
FATHOM_API_URL = "https://api.fathom.ai/external/v1/meetings"

# LOGO ASSET URLS (Replace these URLs with your actual direct image links)
MY_OPSBOX_LOGO_URL = r"MyOpsBox.svg"  # Local file path for My OpsBox logo
STEORA_LOGO_URL = r"Steora.svg"  # Local file path for Steora logo

st.set_page_config(page_title="Discovery Archiver Pro + AI", page_icon=r"Steora-Favicon.png", layout="wide")

# --- BRANDING LAYER: LOGO INTEGRATION ---
logo_col1, logo_col2 = st.columns([1, 1])

with logo_col1:
    # Renders the primary My OpsBox brand logo left-aligned
    try:
        st.image(MY_OPSBOX_LOGO_URL, width=220)
    except Exception:
        # Graceful fallback if the asset URL drops or changes down the road
        st.markdown("### **My OpsBox**")

with logo_col2:
    # Renders the integrated local Steora technology SVG logo right-aligned
    try:
        import os
        
        if os.path.exists(STEORA_LOGO_URL):
            # Read the raw SVG data string from your local drive
            with open(STEORA_LOGO_URL, "r", encoding="utf-8") as svg_file:
                svg_data = svg_file.read()
            
            # Inject the raw SVG directly inside a right-aligned flexbox wrapper
            # We enforce a maximum width of 140px to mirror your original layout constraints
            st.markdown(
                f'<div style="display: flex; justify-content: flex-end;"><div style="width: 140px;">{svg_data}</div></div>', 
                unsafe_allow_html=True
            )
        else:
            # Fallback if the path cannot be found or resolved on this machine
            st.markdown("<p style='text-align: right;'><b>Steora Enabled</b></p>", unsafe_allow_html=True)
            
    except Exception:
        # Graceful error suppression fallback to ensure the UI platform never crashes
        st.markdown("<p style='text-align: right;'><b>Steora Enabled</b></p>", unsafe_allow_html=True)

st.markdown("---")
st.title("🚀 My OpsBox Intelligent Pipeline Engine")
st.markdown("Sync Fathom meetings, run deep-dive strategic analysis via GPT-4o, and generate client-ready proposals instantly.")

# --- SIDEBAR CONFIGURATION ---
st.sidebar.header("⚙️ API Credentials")
fathom_key = st.sidebar.text_input("Fathom API Key", type="password")
openai_key = st.sidebar.text_input("OpenAI API Key", type="password")

st.sidebar.subheader("📅 Sync Parameters")
time_frame = st.sidebar.selectbox("Lookback Window", ["Today", "Past 7 Days", "Past 30 Days", "All Time"])

if not fathom_key or not openai_key:
    st.info("💡 Please input your Fathom and OpenAI API Keys in the sidebar to activate the automated workspace pipeline.")
    st.stop()

# Initialize clients and headers
headers = {"X-Api-Key": fathom_key}
ai_client = OpenAI(api_key=openai_key)

# Calculate lookback constraints
created_after_param = None
if time_frame == "Today":
    created_after_param = (datetime.utcnow() - timedelta(days=1)).strftime("%Y-%m-%dT%H:%M:%SZ")
elif time_frame == "Past 7 Days":
    created_after_param = (datetime.utcnow() - timedelta(days=7)).strftime("%Y-%m-%dT%H:%M:%SZ")
elif time_frame == "Past 30 Days":
    created_after_param = (datetime.utcnow() - timedelta(days=30)).strftime("%Y-%m-%dT%H:%M:%SZ")

# --- DOCUMENT ENGINE: PURE PYTHON WORD & PDF EXPORTERS ---
def text_to_docx_buffer(text_content, title_text):
    doc = Document()
    
    # --- BRAND PALETTE DEFINITIONS (EXACT CLIENT SAMPLE MATCH) ---
    HEX_PRIMARY = "6A398E"                    # Corporate Accent Purple (Hex String)
    COLOR_PRIMARY = RGBColor(106, 57, 142)    # Corporate Accent Purple (RGB)
    COLOR_SECONDARY = RGBColor(115, 115, 115) # Muted Slate Gray 
    COLOR_TEXT = RGBColor(60, 60, 60)         # Charcoal Body Text
    HEX_LIGHT_BG = "F7F8FA"                   # Zebra Striping Light Gray hex
    
    # Configure precise 1-inch standard executive margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- XML HELPER FUNCTIONS FOR BACKGROUND SHADING & THIN BORDERS ---
    def set_cell_background(cell, fill_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_table_borders(table):
        tblPr = table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'D3D3D3') # Clean, thin muted gray border
            tblBorders.append(border)
        tblPr.append(tblBorders)

    # =========================================================================
    # PART A: FIXED TEMPLATE BRAND HEADER (TWO-COLUMN CARD)
    # =========================================================================
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.rows[0].cells[0].width = Inches(3.8)
    header_table.rows[0].cells[1].width = Inches(2.7)
    
    # Left Column: Primary Logo Asset Positioning
    left_cell = header_table.rows[0].cells[0]
    left_p = left_cell.paragraphs[0]
    left_p.paragraph_format.space_after = Pt(2)
    
    MY_OPSBOX_LOGO_PATH = r"MyOpsBox.png"
    import os
    if os.path.exists(MY_OPSBOX_LOGO_PATH):
        try:
            left_p.add_run().add_picture(MY_OPSBOX_LOGO_PATH, width=Inches(2.4))
        except Exception:
            brand_run = left_p.add_run("My OpsBox®")
            brand_run.font.name = 'Arial'; brand_run.font.size = Pt(24); brand_run.font.bold = True; brand_run.font.color.rgb = COLOR_PRIMARY
    else:
        brand_run = left_p.add_run("My OpsBox®")
        brand_run.font.name = 'Arial'; brand_run.font.size = Pt(24); brand_run.font.bold = True; brand_run.font.color.rgb = COLOR_PRIMARY
        
    sub_p = left_cell.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(4)
    sub_run = sub_p.add_run("Operational Assessment & Strategic Analysis")
    sub_run.font.name = 'Arial'; sub_run.font.size = Pt(10.5); sub_run.font.italic = True; sub_run.font.color.rgb = COLOR_SECONDARY
    
    # Right Column: Styled Administrative Contact Sidebar
    right_cell = header_table.rows[0].cells[1]
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    tcPr = right_cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="0" w:color="{HEX_PRIMARY}"/></w:tcBorders>')
    tcPr.append(borders)
    set_cell_margins(right_cell, top=60, bottom=60, left=180, right=60)
    
    for r_text, r_bold, r_color, r_size in [
        ("Prepare by:\n", True, COLOR_SECONDARY, 9.5),
        ("Patti Zapparolli\n", True, COLOR_TEXT, 10),
        ("patti@myopsbox.com\n", False, COLOR_PRIMARY, 9.5),
        ("727-919-7323", False, COLOR_SECONDARY, 9.5)
    ]:
        run = right_p.add_run(r_text)
        run.font.name = 'Arial'; run.font.bold = r_bold; run.font.color.rgb = r_color; run.font.size = Pt(r_size)

    # Accent Divider Break
    sep_p = doc.add_paragraph()
    sep_p.paragraph_format.space_before = Pt(12)
    sep_p.paragraph_format.space_after = Pt(16)
    sep_p.add_run("―" * 60).font.color.rgb = RGBColor(210, 214, 219)

    # =========================================================================
    # PART B: STYLED MARKDOWN INTERPRETATION ENGINE
    # =========================================================================
    lines = text_content.split('\n')
    idx = 0
    
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or line.startswith("---") or line.startswith("___") or line.startswith(":---"):
            idx += 1
            continue

        # 1. PROCESS PIPED DATA MATRICES AND METADATA GRIDS
        if line.startswith("|"):
            table_rows = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                row_raw = lines[idx].strip()
                if not (":---" in row_raw or "---:" in row_raw):
                    row_cells = [c.strip() for c in row_raw.split("|")[1:-1]]
                    if row_cells:
                        table_rows.append(row_cells)
                idx += 1
                
            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                grid_table = doc.add_table(rows=0, cols=max_cols)
                grid_table.style = 'Normal Table'
                set_table_borders(grid_table)
                
                # Check if this table has 4 columns (signaling the top Client Metadata Information box)
                is_metadata_card = (max_cols == 4)
                
                for r_num, row_data in enumerate(table_rows):
                    row_cells = grid_table.add_row().cells
                    
                    # Data metrics tables get full corporate background fills for headers
                    is_header_row = (r_num == 0 and not is_metadata_card)
                    
                    for c_num, val in enumerate(row_data):
                        if c_num < len(row_cells):
                            cell = row_cells[c_num]
                            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                            cell_p = cell.paragraphs[0]
                            cell_p.paragraph_format.space_after = Pt(0)
                            
                            run = cell_p.add_run(val)
                            run.font.name = 'Arial'
                            
                            if is_header_row:
                                set_cell_background(cell, HEX_PRIMARY)
                                run.font.bold = True
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                run.font.size = Pt(9.5)
                                # Color code and bold input form parameters (Cols 0 and 2 of Information Blocks)
                                if is_metadata_card and c_num in [0, 2]:
                                    run.font.bold = True
                                    run.font.color.rgb = COLOR_PRIMARY
                                else:
                                    run.font.color.rgb = COLOR_TEXT
                                    
                                # Smooth gray alternating rows for data tables
                                if not is_metadata_card and r_num % 2 == 1:
                                    set_cell_background(cell, HEX_LIGHT_BG)
                
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                continue

        # 2. RENDER SHADED PRIMARY SECTION BANNER BLOCKS (## Headers)
        if line.startswith("## "):
            clean_title = line.replace("## ", "").replace("**", "").strip()
            
            # Form a single cell table spanning the full width of the text body area
            banner_table = doc.add_table(rows=1, cols=1)
            banner_table.autofit = False
            banner_table.columns[0].width = Inches(6.5)
            
            cell = banner_table.rows[0].cells[0]
            set_cell_background(cell, HEX_PRIMARY) # Corporate Purple Background
            set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(clean_title)
            run.font.name = 'Arial'; run.font.size = Pt(11.5); run.font.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
            
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)
            idx += 1
            continue

        # 3. SUBSECTION HEADINGS (### Headers)
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(line.replace("### ", "").replace("**", "").strip())
            run.font.name = 'Arial'; run.font.size = Pt(11); run.font.bold = True; run.font.color.rgb = COLOR_PRIMARY
            idx += 1
            continue

        # 4. STYLIZED BULLET LIST GENERATOR
        if line.startswith("*") or line.startswith("-") or line.startswith("•"):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            
            clean_line = line.lstrip("*•- ").strip()
            
            # Format inline bold keys if present (e.g., "* Target Parameter: details")
            if " : " in clean_line or "**" in clean_line:
                clean_line = clean_line.replace("**", "")
                splitter = " : " if " : " in clean_line else ":"
                parts = clean_line.split(splitter, 1)
                
                r_bold = p.add_run(parts[0] + splitter)
                r_bold.font.name = 'Arial'; r_bold.font.size = Pt(10.5); r_bold.font.bold = True; r_bold.font.color.rgb = COLOR_PRIMARY
                
                if len(parts) > 1:
                    r_body = p.add_run(parts[1])
                    r_body.font.name = 'Arial'; r_body.font.size = Pt(10.5); r_body.font.color.rgb = COLOR_TEXT
            else:
                run = p.add_run(clean_line)
                run.font.name = 'Arial'; run.font.size = Pt(10.5); run.font.color.rgb = COLOR_TEXT
                
            idx += 1
            continue

        # 5. GENERAL PROSE / NARRATIVE PARAGRAPHS
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        
        clean_text = line.replace("**", "").strip()
        
        # Check if the line contains a descriptor colon to dynamically style it
        if ":" in clean_text and len(clean_text.split(":")[0]) < 30 and not clean_text.startswith("http"):
            parts = clean_text.split(":", 1)
            r_bold = p.add_run(parts[0] + ":")
            r_bold.font.name = 'Arial'; r_bold.font.size = Pt(10.5); r_bold.font.bold = True; r_bold.font.color.rgb = COLOR_PRIMARY
            
            r_body = p.add_run(parts[1])
            r_body.font.name = 'Arial'; r_body.font.size = Pt(10.5); r_body.font.color.rgb = COLOR_TEXT
        else:
            run = p.add_run(clean_text)
            run.font.name = 'Arial'; run.font.size = Pt(10.5); run.font.color.rgb = COLOR_TEXT
            
        idx += 1

    # Return clean binary stream back to the export interface
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

class PurePDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 9)
        self.set_text_color(120, 130, 140)
        self.cell(0, 10, "My OpsBox AI Operational Report", border=0, ln=1, align="R")
        self.ln(3)
    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

def text_to_pdf_buffer(text_content, title_text):
    pdf = PurePDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Title Document Layout
    pdf.set_font("Arial", "B", 18)
    pdf.set_text_color(26, 54, 93)
    pdf.multi_cell(0, 10, txt=title_text)
    pdf.ln(5)
    
    pdf.set_font("Arial", "", 10)
    pdf.set_text_color(60, 60, 60)
    
    for line in text_content.split('\n'):
        if not line.strip():
            continue
            
        # FIX 1: Filter out structural markdown table formatting markers that break FPDF line layout
        if "| :---" in line or "| :---" in line or line.strip() == "|":
            continue
            
        # FIX 2: Safely convert structural markdown divider lines into clean, native line breaks
        if line.strip().startswith("---") or "―――" in line:
            pdf.ln(2)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
            pdf.ln(3)
            continue
            
        # Handle Section Headers (## and ###)
        if line.startswith("##") or line.startswith("###"):
            pdf.ln(2)
            pdf.set_font("Arial", "B", 12)
            pdf.set_text_color(26, 54, 93)
            clean = line.replace("#", "").replace("*", "").strip()
            pdf.multi_cell(0, 6, txt=clean.encode('latin-1', 'ignore').decode('latin-1'))
            pdf.set_font("Arial", "", 10)
            pdf.set_text_color(60, 60, 60)
        else:
            # Clean out common markdown bold highlights within raw lines for standard formatting safety
            clean_line = line.replace("**", "").replace("__", "").strip()
            
            # FIX 3: Safety wrapper to protect the multi_cell execution footprint from edge case character blowouts
            try:
                safe_line = clean_line.encode('latin-1', 'ignore').decode('latin-1')
                pdf.multi_cell(0, 5, txt=safe_line)
                pdf.ln(1)
            except Exception:
                # Fallback to prevent app crash if any aberrant string slips through
                continue
            
    return io.BytesIO(pdf.output())

# --- CENTRAL PIPELINE MECHANICS ---
if st.button("🔄 Sync Active Pipeline", type="primary"):
    with st.spinner("Connecting to Fathom Cloud Gateways..."):
        try:
            params = {"include_transcript": "true", "calendar_invitees_domains_type": "all"}
            if created_after_param:
                params["created_after"] = created_after_param
                
            response = requests.get(FATHOM_API_URL, headers=headers, params=params)
            if response.status_code == 200:
                st.session_state["meetings_list"] = response.json().get("items", [])
                st.success(f"Synced {len(st.session_state['meetings_list'])} calls from Fathom stream.")
            else:
                st.error(f"Fathom error authentication drop: Code {response.status_code}")
        except Exception as e:
            st.error(f"Pipeline failure: {e}")

if "meetings_list" in st.session_state and st.session_state["meetings_list"]:
    meetings = st.session_state["meetings_list"]
    options = [(f"{m.get('created_at', 'TBD')[:10]} — {m.get('title', 'Call')}", idx) for idx, m in enumerate(meetings)]
    
    selected_call = st.selectbox("Select Call Stream for Analysis:", options=options, format_func=lambda x: x[0])
    
    if selected_call is not None:
        target = meetings[selected_call[1]]
        rec_id = target.get("recording_id")
        
        # --- FIX 1: BULLETPROOF DURATION EXTRACTION ---
        # Checks all variations of Fathom API response keys ('duration', 'duration_seconds', etc.)
        raw_duration = target.get('duration') or target.get('duration_seconds') or target.get('length') or 0
        try:
            duration_seconds = int(raw_duration)
        except (ValueError, TypeError):
            duration_seconds = 0
            
        duration_minutes = duration_seconds // 60
        
        # If it still reads 0, check if the call timestamp markers can estimate it
        if duration_minutes == 0 and "transcript" in target and isinstance(target["transcript"], list):
            try:
                # Approximate from the last transcript turn timestamp
                last_turn = target["transcript"][-1]
                ts = last_turn.get("timestamp", "0:00").split(":")
                if len(ts) == 2:
                    duration_minutes = int(ts[0])
                elif len(ts) == 3:
                    duration_minutes = (int(ts[0]) * 60) + int(ts[1])
            except Exception:
                duration_minutes = 15 # Reasonable baseline fallback if completely unexposed
        
        col_m1, col_m2 = st.columns(2)
        current_call_title = target.get("title", "Unknown Call")
        col_m1.metric("Selected Client Call", current_call_title)
        col_m2.metric("Duration Pool", f"{duration_minutes} minutes")
        
        raw_transcript_block = ""
        if rec_id:
            with st.spinner("Downloading raw call data elements..."):
                t_resp = requests.get(f"https://api.fathom.ai/external/v1/recordings/{rec_id}/transcript", headers=headers)
                if t_resp.status_code == 200:
                    turns = t_resp.json().get("transcript", [])
                    raw_transcript_block = "\n".join([f"[{t.get('timestamp', '00:00')}] {t.get('speaker', {}).get('display_name', 'Speaker')}: {t.get('text', '')}" for t in turns])
                else:
                    raw_transcript_block = str(target.get("transcript", ""))

        if not raw_transcript_block.strip() or len(raw_transcript_block) < 20:
            # Handle if the object transcript is a structural list of turns natively nested
            if isinstance(target.get("transcript"), list):
                raw_transcript_block = "\n".join([f"{t.get('speaker', 'Speaker')}: {t.get('text', '')}" for t in target["transcript"]])
            else:
                raw_transcript_block = str(target.get("transcript", "No transcript payload present."))

        st.markdown("---")
        st.subheader("🧠 Operational Intelligence Center")
        st.markdown("Click below to pass raw transcript blocks through your Mastermind prompt structure directly to ChatGPT.")
        
        if st.button("🔥 Run AI Assessment Analysis", type="secondary"):
            if not raw_transcript_block or len(raw_transcript_block) < 30:
                st.error("No valid transcript data found for this call window. Verify Fathom processing state.")
            else:
                with st.spinner(f"Processing deep analysis matrix via GPT-4o for {current_call_title}..."):
                    try:
                        # --- EXPLICIT PIPELINE CORE PROMPT ---
                        master_prompt = f"""
You are an expert Fractional Chief Operating Officer (Fractional COO) acting on behalf of My OpsBox. 
Your objective is to thoroughly analyze the attached raw Fathom call transcript and generate a comprehensive Operational Assessment document that matches the strict tone, structure, and depth shown in your brand's note sheets.

### CORE OPERATIONAL SOLUTIONS KNOWLEDGE (Managed Services Guide Context):
- Leverage the 'My OpsBox Managed Services Guide' positioning matrix. 
- Ensure recommendations cross-reference solutions to explicit tiers: PreOpsBox (Executive Assistant + Backup VA + Fractional COO at $1,850/mo), SubOpsBox ($7,610/mo adding Bookkeeper & Social Media Support), or standalone Specialist offerings (SOP Development, Tech/Web Maintenance, Call Centers).
- Frame solutions under My OpsBox core brand values: Lead from the front, Built for the long haul, Keep it simple and moving, Stay sharp, stay ready, Crew over ego.

### CRITICAL OUTPUT FORMATTING INSTRUCTIONS FOR THE DOCX PROCESSING ENGINE:
Your response must be returned using clean Markdown syntax so the text parser can construct a highly styled Word document matching the 'Assessment Analysis Client Sample'. Follow these structural patterns exactly:

1. CLIENT CARD META BLOCK:
Begin your response immediately with a 4-column structured Markdown metadata table capturing discovered details. Use exact pipe layout filters:
| Business Name | [Insert Discovered Company Name] | Assessment Date | {datetime.now().strftime('%m/%d/%y')} |
| Website | [Insert Client Web Domain] | Email | [Insert Principal Email] |
| Business Type | [Insert Offer / Service Model] | Business Industry | [Insert Industry Segment] |
| Number of full-time employees | [ headcount ] | Annual revenue range | [Estimated run rate or N/A] |

2. PRIMARY HEADINGS:
Every primary block section must explicitly start with a Markdown H2 ('## ') header. Do not output naked text or standalone numbering rules for these:
## 1. Executive Summary & Business Context
## 2. Key Operational Challenges & Gaps
## 3. Strategic 12-Month Objectives
## 4. Proposed Service Package Strategy
## 5. Structured 30-60-90 Day Transformation Timeline Plan
## 6. Scorecard Matrix Projections

3. SUBSECTION TITLE ELEMENTS:
Format subcategories or segment wrappers explicitly with a Markdown H3 ('### ') tag.

4. TARGET SCORECARD DATA MATRIX:
Render Section 6 strictly as a valid, multi-column Markdown data table matching the sample projections grid format:
| Metric | Current State | 30 Days | 60 Days | 90 Days | 12 Months |
| --- | --- | --- | --- | --- | --- |
| [e.g., Client Count / Work Capacity] | [Value] | [Value] | [Value] | [Value] | [Value] |

5. LIST ENHANCEMENTS:
For key breakdowns and item lists, format each point using standard asterisks followed by colons for parameter highlighting (e.g., '* Bandwidth Bottlenecks: Explanatory description text'). Avoid raw comma-separated value (CSV) strings.

### LIVE CALL TRANSCRIPT TEXT TO PROCESS AND ANALYZE:
{raw_transcript_block}
"""
                        ai_response = ai_client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {"role": "system", "content": "You are a world-class Fractional COO and master systems operations analyst specializing in corporate scaling, workflow architecture, and business metrics mapping."},
                                {"role": "user", "content": master_prompt}
                            ],
                            temperature=0.25
                        )
                        
                        # Save the freshly generated analysis safely in session state
                        st.session_state["compiled_analysis"] = ai_response.choices[0].message.content
                        st.success("AI Strategic Assessment Generated Flawlessly From Live Transcript!")
                        st.rerun() # Clear old cached data streams
                    
                    except Exception as ai_err:
                        st.error(f"OpenAI Core Engine Exception: {ai_err}")
                        
        if "compiled_analysis" in st.session_state:
            st.markdown("### 📋 Generated Strategic Analysis")
            st.markdown(st.session_state["compiled_analysis"])
            
            st.markdown("### 📥 Asset Export Options")
            safe_name = "".join([c if c.isalnum() else "_" for c in target.get('title', 'Analysis')])
            
            col_ex1, col_ex2 = st.columns(2)
            with col_ex1:
                if st.button("🛠️ Export Analysis as DOCX"):
                    docx_buf = text_to_docx_buffer(st.session_state["compiled_analysis"], target.get('title', 'Assessment'))
                    st.download_button(
                        label="💾 Save Word Document",
                        data=docx_buf,
                        file_name=f"AI_Assessment_{safe_name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            with col_ex2:
                if st.button("⚖️ Export Analysis as PDF"):
                    pdf_buf = text_to_pdf_buffer(st.session_state["compiled_analysis"], target.get('title', 'Assessment'))
                    st.download_button(
                        label="💾 Save Client PDF",
                        data=pdf_buf,
                        file_name=f"AI_Assessment_{safe_name}.pdf",
                        mime="application/pdf"
                    )
